import os
import pickle
import pdfplumber
import docx
import pandas as pd
import joblib
from django.conf import settings

# Paths to models
BASE_MODEL_DIR = os.path.join(settings.BASE_DIR, 'resume_data_ml_models', 'Models')
MODEL_PATH = os.path.join(BASE_MODEL_DIR, 'best_model_Logistic Regression.pkl')
FEATURES_PATH = os.path.join(BASE_MODEL_DIR, 'features.pkl')
ENCODER_PATH = os.path.join(BASE_MODEL_DIR, 'encoder_classes.pkl')

def extract_text_from_file(file_path):
    """Extracts text from PDF or DOCX files."""
    text = ""
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    return text

def get_page_count(file):
    """Counts pages in a PDF or DOCX file."""
    try:
        if file.name.endswith(".pdf"):
            with pdfplumber.open(file) as pdf:
                return len(pdf.pages)
        elif file.name.endswith(".docx"):
            # DOCX heuristic: 500 words ~ 1 page
            doc = docx.Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            word_count = len(text.split())
            return (word_count // 500) + 1
        return 0
    except Exception:
        return 0

def extract_candidate_name(text, filename):
    """Extracts candidate name from the top of the resume text or falls back to filename."""
    lines = text.split("\n")
    candidate_name = ""

    for line in lines[:10]:
        line = line.strip()
        if len(line) > 3 and len(line) < 40:
            if not any(word in line.lower() for word in ["resume", "cv", "email", "phone", "contact", "experience", "education"]):
                if any(c.isalpha() for c in line):
                    candidate_name = line
                    break

    if candidate_name == "":
        candidate_name = os.path.splitext(filename)[0].replace(".txt", "").replace("_", " ").title()
        
    return candidate_name

def get_rule_based_matches(skills_text, roles_queryset):
    """Calculates rule-based matching for a given set of skills."""
    candidate_skills = [s.strip().lower() for s in skills_text.replace('\n', ',').replace('/', ',').split(",") if s.strip()]
    results = []
    
    for role in roles_queryset:
        role_skills = [s.strip().lower() for s in role.required_skills.split(",") if s.strip()]
        if not role_skills:
            continue
            
        role_skills_set = set(role_skills)
        candidate_skills_set = set(candidate_skills)
        
        matched_skills = sorted(list(candidate_skills_set & role_skills_set))
        missing_skills = sorted(list(role_skills_set - candidate_skills_set))
        
        matched_count = len(matched_skills)
        total = len(role_skills_set)
        match_percentage = round((matched_count / total) * 100, 2) if total > 0 else 0
        
        if match_percentage >= 70:
            fit_category = "Strong Fit"
        elif match_percentage >= 40:
            fit_category = "Moderate Fit"
        else:
            fit_category = "Low Fit"
            
        results.append({
            'role': role,
            'match_percentage': match_percentage,
            'fit_category': fit_category,
            'matched_skills': matched_skills,
            'missing_skills': missing_skills
        })
    
    return results

def get_ml_predictions(extracted_text):
    """Predicts top 3 roles using the trained ML model with one-hot encoding."""
    try:
        if not all(os.path.exists(p) for p in [MODEL_PATH, FEATURES_PATH, ENCODER_PATH]):
            return []

        model = joblib.load(MODEL_PATH)
        with open(FEATURES_PATH, 'rb') as f:
            skills_tools = pickle.load(f)
        with open(ENCODER_PATH, 'rb') as f:
            roles = pickle.load(f)

        candidate_skills = [s.strip().lower() for s in extracted_text.replace('\n', ',').replace('/', ',').split(",") if s.strip()]
        
        X_vec = []
        for skill in skills_tools:
            X_vec.append(1 if skill.lower() in candidate_skills else 0)
            
        prob_matrix = model.predict_proba([X_vec])
        probs = prob_matrix[0]
        role_scores = list(zip(roles, probs * 100))
        role_scores.sort(key=lambda x: x[1], reverse=True)
        top3 = role_scores[:3]
        return top3
    except Exception as e:
        print(f"Error in ML prediction: {e}")
        return []
